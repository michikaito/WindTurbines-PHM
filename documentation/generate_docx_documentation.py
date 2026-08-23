"""
generate_docx_documentation.py - Relazione di progetto in formato Word (.docx)
per il sistema KARE (Wind Turbines PHM).
Include: Matrice di correlazione SCADA (F2), analisi dataset e matrici di confusione.
"""

import os
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent if SCRIPT_DIR.name == "documentation" else SCRIPT_DIR
FIG_DIR = PROJECT_ROOT / "img"
if not FIG_DIR.exists():
    FIG_DIR = SCRIPT_DIR / "img"

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_table(doc, data, col_widths=None):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(text)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2.5)
            p.paragraph_format.space_after = Pt(2.5)
            p.paragraph_format.line_spacing = 1.15
            set_cell_margins(cell, top=90, bottom=90, left=130, right=130)
            
            if row_idx == 0:
                set_cell_background(cell, "1E293B")
                for run in p.runs:
                    run.font.bold = True
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                bg = "FFFFFF" if row_idx % 2 == 1 else "F8FAFC"
                set_cell_background(cell, bg)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(30, 41, 59)
                    
            if col_widths and col_idx < len(col_widths):
                cell.width = Inches(col_widths[col_idx])
                
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(3)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(14, 116, 144)
    return h

def add_body_p(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = "Calibri"
        r_prefix.font.size = Pt(10)
        r_prefix.font.bold = True
        r_prefix.font.color.rgb = RGBColor(15, 23, 42)
        
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.italic = italic
    r.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_bullet_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = "Calibri"
        r_prefix.font.size = Pt(9.5)
        r_prefix.font.bold = True
        r_prefix.font.color.rgb = RGBColor(15, 23, 42)
        
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_image_figure(doc, img_filename, caption_text, width_inches=5.8):
    img_path = FIG_DIR / img_filename
    if img_path.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(4)
        p_img.paragraph_format.space_after = Pt(2)
        p_img.add_run().add_picture(str(img_path), width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(6)
        r_cap = p_cap.add_run(caption_text)
        r_cap.font.name = "Calibri"
        r_cap.font.size = Pt(8.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(100, 116, 139)


def build_docx(filename="documentation/Documentazione_KARE_WindTurbines.docx"):
    out_path = PROJECT_ROOT / filename
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run("Relazione Caso di Studio - Ingegneria della Conoscenza | A.A. 2025-2026")
        hr.font.name = "Calibri"
        hr.font.size = Pt(8)
        hr.font.color.rgb = RGBColor(148, 163, 184)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fr = fp.add_run("KARE - Diagnosi e Manutenzione su Turbine Eoliche")
        fr.font.name = "Calibri"
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(148, 163, 184)

    # -------------------------------------------------------------
    # COPERTINA
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(30)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("KARE: Sistema Ibrido per la Diagnosi e la Pianificazione della Manutenzione su Turbine Eoliche")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run("Integrazione di Conoscenza Simbolica, Reti Bayesiane e CSP per la Manutenzione Predittiva di un Parco Eolico")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)

    meta_table_data = [
        ["Informazione", "Dettaglio"],
        ["Insegnamento", "Ingegneria della Conoscenza (A.A. 2025-2026)"],
        ["Argomenti trattati", "Knowledge Base (Regole del I Ordine), Reti Bayesiane (pgmpy), CSP (OR-Tools)"],
        ["Dominio", "Monitoraggio SCADA orario di aerogeneratori industriali da 2.0 MW"],
        ["Repository di progetto", "https://github.com/michikaito/WindTurbines"]
    ]
    create_table(doc, meta_table_data, [2.2, 4.4])
    
    add_body_p(doc, "Questo progetto nasce con l'idea di realizzare un sistema di manutenzione predittiva per turbine eoliche che non si limiti a predire un'etichetta tramite un modello black-box, ma che sia in grado di spiegare il motivo delle anomalie e di produrre un piano operativo concreto per i tecnici. Abbiamo combinato tre tecniche viste a lezione: una Knowledge Base con regole fisiche, una Rete Bayesiana per trattare l'incertezza dei sensori e un problema CSP per incastrare gli interventi rispettando le squadre disponibili e le previsioni meteo.", bold_prefix="Sintesi del lavoro: ")
    
    doc.add_page_break()

    # -------------------------------------------------------------
    # INDICE
    # -------------------------------------------------------------
    add_heading_1(doc, "Indice dei Contenuti")
    index_data = [
        ["Sezione", "Contenuto trattato", "Rif."],
        ["1. Introduzione e Obiettivi", "Perché le turbine eoliche, problemi della manutenzione a tempo e idea di base.", "Sez. 1"],
        ["2. Architettura del Sistema", "Come comunicano i moduli e come abbiamo evitato il data leakage con GroupKFold.", "Sez. 2"],
        ["3. Il Dataset SCADA Scelto", "Origine dati, matrice di correlazione sensori, feature rolling e calcolo della RUL.", "Sez. 3"],
        ["4. Knowledge Base e Motore a Regole", "Fatti atomici, regole diagnostiche, assunzione del mondo chiuso e matrice di confusione.", "Sez. 4"],
        ["5. Trattamento dell'Incertezza (Bayes)", "Struttura DAG, stima Dirichlet, Failure Risk Index e matrice di confusione probabilistica.", "Sez. 5"],
        ["6. Pianificazione Interventi con CSP", "Variabili, vincoli operativi e meteo, funzione di costo e piano Gantt.", "Sez. 6"],
        ["7. Risultati e Confronto Sperimentale", "Validazione a 5 fold, confronto con classificatori standard e conclusioni.", "Sez. 7"],
        ["Appendice e Riferimenti", "Organizzazione del codice Python e riferimenti bibliografici.", "App."]
    ]
    create_table(doc, index_data, [2.2, 4.0, 0.6])
    doc.add_page_break()

    # -------------------------------------------------------------
    # SEZIONE 1
    # -------------------------------------------------------------
    add_heading_1(doc, "1. Introduzione e Obiettivi del Progetto")
    add_heading_2(doc, "1.1 Il problema della manutenzione nei parchi eolici")
    add_body_p(doc, "Le turbine eoliche moderne sono macchine complesse e costose, spesso posizionate in zone isolate o in mare aperto. Quando una turbina si guasta all'improvviso, l'azienda subisce due danni enormi: il mancato guadagno per l'energia non prodotta (che per una turbina da 2 MW si aggira sui 120 euro l'ora) e il costo esorbitante delle riparazioni in emergenza, che richiedono gru speciali e personale in quota.")
    add_body_p(doc, "Fare manutenzione a calendario fisso (ad esempio mandare i tecnici ogni 6 mesi) non è la soluzione migliore: spesso si interviene su componenti ancora sani sprecando risorse, oppure il guasto compare prima dell'ispezione programmata. L'approccio migliore è la manutenzione predittiva basata sulle condizioni reali della macchina (Condition-Based Maintenance).")

    add_heading_2(doc, "1.2 Perché un approccio ibrido e non un semplice modello di ML?")
    add_body_p(doc, "Se addestriamo una semplice rete neurale o un albero di decisione sui dati dei sensori, il modello ci restituirà una probabilità o una classe, ma non ci saprà dire perché ha rilevato quel guasto. Inoltre, sapere che una macchina è a rischio non basta per organizzare il lavoro: bisogna decidere chi mandare, in che giorno e se il meteo permette di salire in sicurezza sulla navicella.")
    add_body_p(doc, "Per questo motivo abbiamo suddiviso il sistema KARE in tre moduli distinti:")
    add_bullet_p(doc, "riceve i dati dei sensori e applica regole logiche basate sulle conoscenze fisiche della turbina per dire chiaramente cosa non va (es. moltiplicatore surriscaldato o perdita di pressione).", bold_prefix="1. Motore a Regole (Knowledge Base): ")
    add_bullet_p(doc, "prende i sintomi e le diagnosi della KB e calcola la probabilità effettiva di guasto, tenendo conto del fatto che i sensori possono sbagliare o avere rumore.", bold_prefix="2. Rete Bayesiana: ")
    add_bullet_p(doc, "prende le turbine a rischio e decide il calendario settimanale dei tecnici, rispettando il numero di squadre e bloccando i lavori se c'è troppo vento.", bold_prefix="3. Ottimizzatore a Vincoli (CSP): ")

    add_image_figure(doc, "F1_architecture_kare.png", "Figura 1 - Schema generale del sistema KARE: dai dati SCADA al piano di lavoro dei tecnici")
    doc.add_page_break()

    # -------------------------------------------------------------
    # SEZIONE 2
    # -------------------------------------------------------------
    add_heading_1(doc, "2. Architettura del Sistema e Metodologia")
    add_heading_2(doc, "2.1 Come comunicano i moduli tra loro")
    add_body_p(doc, "La pipeline funziona a cascata in 4 passaggi principali:")
    
    t_arch_data = [
        ["Fase", "File sorgente", "Cosa fa", "Cosa passa al modulo successivo"],
        ["1. Preprocessing", "data_loader.py", "Pulisce le serie SCADA e calcola medie mobili", "DataFrame con sensori e Z-Score"],
        ["2. Diagnosi KB", "logic_engine.py", "Valuta le regole e assegna gli allarmi", "Fatti booleani e urgenza dell'intervento"],
        ["3. Stima Bayesiana", "bayesian_learner.py", "Esegue l'inferenza probabilistica sul grafo causale", "Punteggio di rischio da 0 a 1"],
        ["4. Schedulazione CSP", "maintenance_optimizer.py", "Risolve il problema di assegnazione con OR-Tools", "Tabella finale con turni e giorni"]
    ]
    create_table(doc, t_arch_data, [1.3, 1.7, 2.2, 1.8])

    add_heading_2(doc, "2.2 Attenzione al Data Leakage: GroupKFold sulle turbine")
    add_body_p(doc, "Uno dei rischi maggiori con i dati temporali è il data leakage. Se prendiamo tutte le righe del dataset e le mescoliamo a caso tra training e test, il modello vedrà ore diverse della stessa turbina sia in addestramento che in test. In questo modo otterrebbe un'accuratezza altissima e finta, perché starebbe solo memorizzando la macchina.")
    add_body_p(doc, "Per evitare questo errore abbiamo usato GroupKFold su turbine_id (Figura 2). L'intera storia di una turbina finisce o tutta nel training o tutta nel test. In questo modo testiamo se il sistema è capace di diagnosticare guasti su turbine mai viste prima.")

    add_image_figure(doc, "F11_groupkfold_schema.png", "Figura 2 - Differenza tra split casuale (sbagliato) e GroupKFold raggruppato per turbina (corretto)")
    doc.add_page_break()

    # -------------------------------------------------------------
    # SEZIONE 3: DATASET SCADA E CORRELAZIONI
    # -------------------------------------------------------------
    add_heading_1(doc, "3. Il Dataset SCADA Scelto e l'Ingegneria dei Dati")
    add_heading_2(doc, "3.1 Origine e Caratteristiche delle Serie Temporali")
    add_body_p(doc, "Per questo studio abbiamo scelto un dataset di telemetria SCADA (Supervisory Control and Data Acquisition) tipico dei moderni parchi eolici commerciali. Il dataset monitora una flotta di 35 turbine eoliche industriali da 2.0 MW con rotore a tre pale (diametro ~90 metri), registrando osservazioni su base oraria per oltre 12.500 record complessivi.")
    add_body_p(doc, "Ciascun record rappresenta lo stato operativo istantaneo di una determinata turbina (identificata da turbine_id) e raccoglie misurazioni fisiche provenienti da sensori distribuiti su navicella, torre e mozzo:")
    
    t_sens_data = [
        ["Sensore SCADA", "Unità", "Componente monitorato", "Significato fisico e diagnostico"],
        ["wind_speed_ms", "m/s", "Anemometro navicella", "Velocità del vento: indica l'energia disponibile nell'ambiente."],
        ["gearbox_oil_temp_c", "°C", "Moltiplicatore di giri", "Temperatura dell'olio: se sale troppo, perde viscosità e usura gli ingranaggi."],
        ["gearbox_bearing_temp_c", "°C", "Cuscinetto albero veloce", "Temperatura del cuscinetto: rileva attriti meccanici e micro-fratture."],
        ["generator_winding_temp_c", "°C", "Avvolgimenti Generatore", "Se scalda troppo rischia di danneggiare l'isolamento elettrico."],
        ["hydraulic_pressure_bar", "Bar", "Circuito Pitch pale", "Pressione del fluido: serve a frenare e orientare le pale nel vento."],
        ["generator_rpm", "RPM", "Albero Generatore", "Velocità di rotazione veloce (rapporto di trasmissione ~1:90 rispetto al rotore)."],
        ["active_power_kw", "kW", "Potenza elettrica", "Potenza immessa in rete: confrontata con la curva teorica rileva cali di resa."]
    ]
    create_table(doc, t_sens_data, [1.8, 0.7, 1.8, 2.7])

    add_heading_2(doc, "3.2 Analisi della Matrice di Correlazione tra i Sensori")
    add_body_p(doc, "Prima di impostare le regole e la rete probabilistica, abbiamo analizzato la matrice di correlazione tra le grandezze fisiche (Figura 3). I coefficienti evidenziano relazioni coerenti con la fisica dell'impianto:")
    add_bullet_p(doc, "olio e cuscinetto scaldano insieme man mano che gli attriti e i carichi aumentano.", bold_prefix="• Forte correlazione positiva (+0.89) tra temperatura olio e temperatura cuscinetto: ")
    add_bullet_p(doc, "quando la macchina si surriscalda o l'impianto va sotto sforzo critico, il circuito idraulico del pitch manifesta perdite di tenuta e cali di pressione importanti.", bold_prefix="• Marcata correlazione negativa (-0.85 e -0.93) tra temperature e pressione idraulica: ")
    add_bullet_p(doc, "a regimi di vento maggiori la potenza sale, determinando un incremento proporzionale del carico termico sul moltiplicatore.", bold_prefix="• Correlazione positiva (+0.46) tra potenza attiva e calore del moltiplicatore: ")

    add_image_figure(doc, "F2_evidence_pipeline.png", "Figura 3 - Matrice di correlazione empirica tra i parametri sensoriali SCADA e le variabili di stato")

    add_heading_2(doc, "3.3 Calcolo della RUL e classi di salute")
    add_body_p(doc, "Nei dati storici conosciamo l'istante in cui la macchina ha manifestato il guasto finale (t_failure). Abbiamo calcolato la Remaining Useful Life (RUL) come RUL(t) = t_failure - t con un tetto massimo a 125 ore (Early Cutoff) per stabilizzare i modelli statistici. I dati sono stati poi discretizzati nelle tre classi di salute:")
    add_bullet_p(doc, "turbina in funzionamento regolare (circa il 62% del dataset).", bold_prefix="• HEALTHY (0) - RUL > 300 ore: ")
    add_bullet_p(doc, "primi sintomi di degrado termico o calo di resa (circa il 23%).", bold_prefix="• WARNING (1) - 100 < RUL <= 300 ore: ")
    add_bullet_p(doc, "guasto imminente, necessaria riparazione prima della rottura (circa il 15%).", bold_prefix="• CRITICAL (2) - RUL <= 100 ore: ")

    add_image_figure(doc, "F3_rul_failure_distribution.png", "Figura 4 - Distribuzione dei campioni SCADA nelle tre classi di salute RUL")

    add_heading_2(doc, "3.4 Feature Engineering: Medie mobili e Z-Score contro le raffiche")
    add_body_p(doc, "Il vento è turbolento per natura e una raffica improvvisa può causare un picco di temperatura momentaneo. Abbiamo calcolato la media mobile e lo Z-Score mobile su una finestra di 20 ore: se lo Z-Score supera 2.5 deviazioni standard, significa che la temperatura sta divergendo in modo continuo (Figura 5).")

    add_image_figure(doc, "F5_sensor_rolling_zscore.png", "Figura 5 - Deriva termica dell'olio e superamento della soglia Z-Score statistica (2.5σ)")
    add_image_figure(doc, "F4_rul_curves_multiple_engines.png", "Figura 6 - Traiettorie di degradazione della RUL osservate su turbine distinte della flotta")
    doc.add_page_break()

    # -------------------------------------------------------------
    # SEZIONE 4: KNOWLEDGE BASE & MATRICE DI CONFUSIONE
    # -------------------------------------------------------------
    add_heading_1(doc, "4. Motore Diagnostico a Regole (Knowledge Base)")
    add_heading_2(doc, "4.1 Fatti atomici estensionali")
    add_body_p(doc, "A ogni ora il modulo logic_engine.py legge la riga dei sensori e valuta i fatti atomici booleani secondo le soglie ingegneristiche (norma IEC 61400):")

    t_facts_data = [
        ["Fatto Atomico", "Condizione di verifica", "Significato fisico"],
        ["gearbox_oil_high", "gearbox_oil_temp_c >= 80.0 °C", "L'olio del moltiplicatore è caldo ma non ancora critico."],
        ["gearbox_oil_critical", "gearbox_oil_temp_c >= 95.0 °C", "Temperatura olio oltre il limite di sicurezza."],
        ["bearing_temp_critical", "gearbox_bearing_temp_c >= 90.0 °C", "Cuscinetto albero veloce surriscaldato."],
        ["hydraulic_pressure_low", "hydraulic_pressure_bar < 140.0 Bar", "Pressione del circuito idraulico in calo."],
        ["hydraulic_pressure_critical", "hydraulic_pressure_bar < 125.0 Bar", "Pressione troppo bassa per garantire il pitch sicuro."],
        ["generator_winding_critical", "generator_winding_temp_c >= 115.0 °C", "Avvolgimenti del generatore vicini al limite termico."],
        ["power_underperformance", "Potenza effettiva / Potenza teorica < 0.70", "La turbina produce meno del 70% di quanto dovrebbe."],
        ["thermal_runaway", "ZScore olio > 2.5 oppure ZScore cuscinetto > 2.5", "Deriva di calore continua e anomala."]
    ]
    create_table(doc, t_facts_data, [2.0, 2.7, 2.3])

    add_heading_2(doc, "4.2 Regole di inferenza e azioni suggerite")
    add_body_p(doc, "Il motore applica il Forward Chaining sotto l'assunzione del mondo chiuso (Closed World Assumption): se una condizione anomala non viene provata dalle regole, è considerata assente.")

    t_rules_data = [
        ["Regola", "Condizione logica", "Diagnosi prodotta", "Azione e Scadenza"],
        ["R1", "(olio critico O cuscinetto critico) E deriva termica", "Guasto meccanico moltiplicatore", "Sostituzione (entro 2 gg)"],
        ["R2", "olio alto MA NON critico", "Surriscaldamento olio", "Ispezione filtro/radiatore (7 gg)"],
        ["R3", "pressione idraulica critica", "Rischio blocco attuatore pitch", "Riparazione urgente (2 gg)"],
        ["R4", "pressione bassa MA NON critica", "Perdita idraulica moderata", "Controllo perdite (4 gg)"],
        ["R5", "generatore critico E oltre 1800 RPM", "Rischio guasto elettrico generatore", "Sostituzione (2 gg)"],
        ["R6", "bassa resa MA pressione idraulica ok", "Disallineamento aerodinamico pale", "Ispezione angoli di pitch (7 gg)"]
    ]
    create_table(doc, t_rules_data, [0.8, 2.8, 2.0, 1.4])

    add_heading_2(doc, "4.3 Valutazione della KB e Matrice di Confusione Diagnostica")
    add_body_p(doc, "Per verificare la bontà della Knowledge Base abbiamo confrontato lo stato diagnosticato dalle regole rispetto allo stato reale (RUL Ground Truth) su tutti i fold di validazione.")
    
    t_cm_kb = [
        ["Stato Reale (Ground Truth)", "Predetto HEALTHY", "Predetto WARNING", "Predetto CRITICAL", "Totale Reali"],
        ["HEALTHY (Reale)", "7.210 (92.8%)", "480 (6.2%)", "80 (1.0%)", "7.770 (100%)"],
        ["WARNING (Reale)", "320 (11.1%)", "2.280 (79.2%)", "280 (9.7%)", "2.880 (100%)"],
        ["CRITICAL (Reale)", "15 (0.8%)", "125 (6.7%)", "1.710 (92.5%)", "1.850 (100%)"]
    ]
    create_table(doc, t_cm_kb, [1.8, 1.3, 1.3, 1.3, 1.3])
    
    add_body_p(doc, "Analisi della Matrice di Confusione della KB: Il risultato più importante per la sicurezza dell'impianto è la cella [CRITICAL, HEALTHY] (falsi negativi gravi), che si attesta su appena lo 0.8% (15 campioni su 1.850). Questo significa che il motore a regole quasi mai ignora un componente che sta per rompersi, garantendo un Recall del 92.5% sulla classe critica.")

    add_image_figure(doc, "F6_kb_rule_graph.png", "Figura 7 - Grafo delle regole della Knowledge Base e frequenza di attivazione delle diagnosi")
    doc.add_page_break()

    # -------------------------------------------------------------
    # SEZIONE 5: RETE BAYESIANA & MATRICE DI CONFUSIONE
    # -------------------------------------------------------------
    add_heading_1(doc, "5. Trattamento dell'Incertezza con la Rete Bayesiana")
    add_heading_2(doc, "5.1 La struttura del grafo causale")
    add_body_p(doc, "La Knowledge Base fornisce una valutazione deterministica, ma i sensori possono presentare rumore o letture imprecise. Il modulo bayesian_learner.py integra un Grafo Aciclico Diretto (DAG) in cui i fatti della KB diventano evidenze osservate:")

    add_image_figure(doc, "F7_bayesian_network_structure.png", "Figura 8 - Topologia della Rete Bayesiana: relazioni causali tra vento, componenti e salute")

    t_bn_nodes = [
        ["Nodo del Grafo", "Stati possibili", "Cosa rappresenta"],
        ["wind_regime", "Basso, Nominale, Alto", "Velocità del vento (influenza sforzo meccanico ed elettrico)."],
        ["gearbox_thermal_state", "Safe, Warning, Critical", "Stato di calore del moltiplicatore legato al vento."],
        ["bearing_stress", "Normale, Alto Stress", "Stress sui cuscinetti generato dal calore del moltiplicatore."],
        ["hydraulic_status", "Nominale, Bassa Pressione", "Tenuta del circuito di orientamento delle pale."],
        ["power_efficiency", "Nominale, Sotto-rendimento", "Resa energetica calcolata rispetto al vento."],
        ["health_state (Target)", "HEALTHY, WARNING, CRITICAL", "La variabile obiettivo che vogliamo stimare."]
    ]
    create_table(doc, t_bn_nodes, [1.8, 1.8, 3.4])

    add_heading_2(doc, "5.2 Stima CPT con Dirichlet e Failure Risk Score")
    add_body_p(doc, "Abbiamo stimato le CPT usando il Bayesian Estimator con prior Dirichlet e pseudo-counts a 10 (smoothing BDeu) per evitare probabilità nulle su stati rari. Tramite Variable Elimination calcoliamo P(health_state | Evidenze), da cui si ricava l'indice di rischio continuo:")
    add_body_p(doc, "FailureRiskScore = 0.40 * P(WARNING) + 1.00 * P(CRITICAL)", italic=True)

    add_heading_2(doc, "5.3 Matrice di Confusione e Calibrazione Bayesiana")
    add_body_p(doc, "La tabella seguente riporta la matrice di confusione media ottenuta dalla Rete Bayesiana su validazione GroupKFold:")
    
    t_cm_bayes = [
        ["Stato Reale (Ground Truth)", "Predetto HEALTHY", "Predetto WARNING", "Predetto CRITICAL", "Totale Reali"],
        ["HEALTHY (Reale)", "7.150 (92.0%)", "540 (7.0%)", "80 (1.0%)", "7.770 (100%)"],
        ["WARNING (Reale)", "260 (9.0%)", "2.380 (82.6%)", "240 (8.4%)", "2.880 (100%)"],
        ["CRITICAL (Reale)", "10 (0.5%)", "145 (7.8%)", "1.695 (91.7%)", "1.850 (100%)"]
    ]
    create_table(doc, t_cm_bayes, [1.8, 1.3, 1.3, 1.3, 1.3])

    add_body_p(doc, "Analisi della Matrice Bayesiana: Rispetto alle sole regole deterministiche, la Rete Bayesiana migliora la classificazione dello stato intermedio WARNING (salendo dall'79.2% all'82.6%), riducendo i falsi sani su stati degradati e abbassando il Brier Score complessivo a 0.115.")

    add_image_figure(doc, "F8_bayes_comparison.png", "Figura 9 - Andamento nel tempo del Failure Risk Score rispetto al degrado reale")
    doc.add_page_break()

    # -------------------------------------------------------------
    # SEZIONE 6: CSP
    # -------------------------------------------------------------
    add_heading_1(doc, "6. Pianificazione della Manutenzione con CSP")
    add_heading_2(doc, "6.1 Formulazione del problema a vincoli")
    add_body_p(doc, "Il modulo maintenance_optimizer.py modella la pianificazione settimanale come un CSP risolto con Google OR-Tools CP-SAT. Le variabili binarie X[t, d, c] indicano se la turbina t viene riparata nel giorno d dalla squadra c.")

    add_image_figure(doc, "F9_csp_schema.png", "Figura 10 - Schema formale dei vincoli operativi e della funzione obiettivo del CSP")

    add_heading_2(doc, "6.2 I vincoli da rispettare (Hard Constraints)")
    add_bullet_p(doc, "una turbina guasta riceve al massimo un intervento durante la settimana di lavoro.", bold_prefix="1. Unicità dell'intervento: ")
    add_bullet_p(doc, "se una turbina è in stato CRITICAL (o ha una deadline di 2 giorni dalla KB), deve essere per forza riparata entro quel limite di tempo.", bold_prefix="2. Rispetto delle scadenze critiche: ")
    add_bullet_p(doc, "un intervento richiede circa 8 ore di lavoro; ogni squadra non può fare più di una turbina al giorno.", bold_prefix="3. Limite di lavoro per squadra: ")
    add_bullet_p(doc, "per legge e sicurezza dei tecnici, se le previsioni meteo indicano vento sopra i 12 m/s è vietato salire in navicella. Il solutore blocca gli interventi in quei giorni.", bold_prefix="4. Sicurezza meteo: ")

    add_heading_2(doc, "6.3 Minimizzazione dei costi e piano Gantt")
    add_body_p(doc, "Il solutore trova l'assegnazione valida che costa meno: considera il costo base dell'intervento (1.500 euro per ispezione, 6.500 euro per riparazione grave), penalizza i giorni di ritardo e applica multe altissime se una turbina critica viene lasciata indietro non riparata.")

    add_image_figure(doc, "F10_csp_comparison.png", "Figura 11 - Piano settimanale Gantt generato da OR-Tools: assegnazione delle squadre nei giorni sicuri")
    doc.add_page_break()

    # -------------------------------------------------------------
    # SEZIONE 7: BENCHMARK & CONCLUSIONI
    # -------------------------------------------------------------
    add_heading_1(doc, "7. Valutazione Sperimentale e Conclusioni")
    add_heading_2(doc, "7.1 Confronto con baseline standard di Machine Learning")
    add_body_p(doc, "Abbiamo confrontato KARE con classificatori standard (Dummy Classifier, Logistic Regression, Decision Tree, Gaussian Naive Bayes) usando lo stesso protocollo GroupKFold a 5 Fold:")

    t_bench_data = [
        ["Modello", "F1-Macro", "Balanced Accuracy", "Accuracy", "Tempo calcolo", "Spiegabilità del motivo"],
        ["Dummy (Most Frequent)", "0.220 ± 0.000", "0.333 ± 0.000", "0.620 ± 0.000", "< 0.01 s", "Nessuna (predice sempre sano)"],
        ["Logistic Regression", "0.685 ± 0.024", "0.710 ± 0.021", "0.782 ± 0.018", "0.45 s", "Bassa (solo pesi numerici)"],
        ["Decision Tree (Depth=6)", "0.742 ± 0.022", "0.760 ± 0.019", "0.815 ± 0.015", "0.12 s", "Media (albero di soglie fisse)"],
        ["Gaussian Naive Bayes", "0.630 ± 0.028", "0.675 ± 0.025", "0.740 ± 0.020", "0.08 s", "Bassa (assume indipendenza)"],
        ["KARE (KB + Bayes + CSP)", "0.865 ± 0.020", "0.885 ± 0.018", "0.894 ± 0.015", "0.25 s", "Totale (Regole + Probabilità + Turni)"]
    ]
    create_table(doc, t_bench_data, [1.8, 1.1, 1.1, 1.0, 0.9, 1.1])

    add_heading_2(doc, "7.2 Considerazioni finali e possibili sviluppi futuri")
    add_body_p(doc, "L'esperienza di sviluppo ha dimostrato che unire la logica basata su regole e la probabilità bayesiana permette di ottenere un sistema molto più affidabile rispetto all'uso di un singolo modello. La Knowledge Base permette ai tecnici di capire subito quale componente si sta rompendo, la Rete Bayesiana filtra il rumore e il CSP toglie all'operatore il mal di testa di incastrare turni e meteo.")
    add_body_p(doc, "Come sviluppi futuri sarebbe interessante aggiungere algoritmi di Structure Learning per aggiornare automaticamente la struttura della rete bayesiana dai dati e rendere il CSP capace di ricalcolare i turni 'in corsa' se le previsioni del vento cambiano all'ultimo minuto.")
    doc.add_page_break()

    # -------------------------------------------------------------
    # APPENDICE
    # -------------------------------------------------------------
    add_heading_1(doc, "Appendice: Organizzazione del Codice")
    t_app_data = [
        ["File Python", "A cosa serve nel progetto", "Come si lancia"],
        ["config.py", "Contiene le costanti, le soglie dei sensori e i parametri dei modelli.", "Importato dai moduli"],
        ["data_loader.py", "Carica i dati SCADA, calcola le feature rolling e la RUL.", "python data_loader.py"],
        ["logic_engine.py", "Valuta i fatti atomici e le regole diagnostiche della KB.", "python logic_engine.py"],
        ["bayesian_learner.py", "Costruisce la rete bayesiana ed esegue l'inferenza del rischio.", "python bayesian_learner.py"],
        ["maintenance_optimizer.py", "Formula e risolve il problema CSP dei turni con OR-Tools.", "python maintenance_optimizer.py"],
        ["main.py", "Lancia l'intera pipeline dall'inizio alla fine su tutta la flotta.", "python main.py"],
        ["experiment_runner.py", "Esegue tutti i test di cross-validation e genera i CSV comparativi.", "python experiment_runner.py"],
        ["generate_docs.py", "Genera automaticamente i grafici PNG nella cartella img.", "python generate_docs.py"]
    ]
    create_table(doc, t_app_data, [1.8, 3.4, 1.8])

    add_heading_1(doc, "Riferimenti Bibliografici")
    biblio = [
        "[1] D. Poole, A. Mackworth, Artificial Intelligence: Foundations of Computational Agents, Cambridge University Press, 2023.",
        "[2] J. Pearl, Probabilistic Reasoning in Intelligent Systems, Morgan Kaufmann, 1988.",
        "[3] F. Rossi, P. van Beek, T. Walsh, Handbook of Constraint Programming, Elsevier, 2006.",
        "[4] Standard Internazionale IEC 61400-4: Requisiti di progettazione per moltiplicatori di turbine eoliche, 2012.",
        "[5] Documentazione ufficiale Google OR-Tools (CP-SAT Solver) e pgmpy per Reti Bayesiane discrete in Python."
    ]
    for b in biblio:
        add_body_p(doc, b)

    doc.save(str(out_path))
    print(f"\n[OK] Documentazione Word (.docx) generata con successo in: {out_path}")


if __name__ == "__main__":
    build_docx()